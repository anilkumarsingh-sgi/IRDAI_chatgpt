"""
IRDAI Compliance GPT - Ingestion Layer
PDF / Excel / Word → Text Extraction → Chunking → Embedding → ChromaDB Storage
"""

import os
import logging
from pathlib import Path
from typing import List

import chromadb
from chromadb.config import Settings
import pdfplumber
import openpyxl
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer

logger = logging.getLogger("irdai.ingestion")

# ─── Config ────────────────────────────────────────────────────────────────────
_ON_CLOUD = Path("/mount/src").exists()
_DATA_ROOT = Path("/tmp/irdai_data") if _ON_CLOUD else Path("data")

PDF_DIR        = _DATA_ROOT / "pdfs"
EXCEL_DIR      = _DATA_ROOT / "excel"
WORD_DIR       = _DATA_ROOT / "word"
CHROMA_DIR     = _DATA_ROOT / "chroma_db"
EMBED_MODEL    = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE     = 800
CHUNK_OVERLAP  = 100
COLLECTION_NAME = "irdai_docs"


# ─── Text Extraction ───────────────────────────────────────────────────────────
def extract_text_from_pdf(pdf_path: Path) -> List[dict]:
    """
    Extract text from each page of a PDF.
    Returns list of {"page": int, "text": str, "source": str}
    """
    pages = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                text = text.strip()
                if len(text) > 50:  # skip near-empty pages
                    pages.append({
                        "page":   i,
                        "text":   text,
                        "source": pdf_path.name,
                    })
    except Exception as exc:
        logger.error("Failed to extract %s: %s", pdf_path.name, exc)
    return pages


def extract_text_from_excel(excel_path: Path) -> List[dict]:
    """
    Extract text from all sheets of an Excel file.
    Each sheet becomes one 'page'. Cell values are joined into readable text.
    """
    pages = []
    try:
        wb = openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
        for sheet_idx, sheet_name in enumerate(wb.sheetnames, start=1):
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                cell_vals = [str(c).strip() for c in row if c is not None]
                if cell_vals:
                    rows_text.append(" | ".join(cell_vals))
            text = "\n".join(rows_text).strip()
            if len(text) > 50:
                pages.append({
                    "page":   sheet_idx,
                    "text":   text,
                    "source": excel_path.name,
                })
        wb.close()
    except Exception as exc:
        logger.error("Failed to extract %s: %s", excel_path.name, exc)
    return pages


def extract_text_from_word(word_path: Path) -> List[dict]:
    """
    Extract text from a Word (.docx) file.
    All paragraphs are treated as page 1 (Word docs don't have page metadata).
    """
    pages = []
    try:
        doc = DocxDocument(str(word_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n".join(paragraphs).strip()
        if len(text) > 50:
            pages.append({
                "page":   1,
                "text":   text,
                "source": word_path.name,
            })
    except Exception as exc:
        logger.error("Failed to extract %s: %s", word_path.name, exc)
    return pages


# ─── Chunking ──────────────────────────────────────────────────────────────────
def chunk_pages(pages: List[dict]) -> List[dict]:
    """
    Split page texts into smaller overlapping chunks.
    Preserves source and page metadata.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = []
    for page in pages:
        splits = splitter.split_text(page["text"])
        for idx, chunk in enumerate(splits):
            chunks.append({
                "text":   chunk,
                "source": page["source"],
                "page":   page["page"],
                "chunk":  idx,
            })
    return chunks


# ─── Embedding & Vector Store ──────────────────────────────────────────────────
_embed_model: SentenceTransformer | None = None

def get_embed_model() -> SentenceTransformer:
    """Lazy-load embedding model (cached)."""
    global _embed_model
    if _embed_model is None:
        logger.info("Loading embedding model: %s", EMBED_MODEL)
        _embed_model = SentenceTransformer(EMBED_MODEL)
    return _embed_model


def get_chroma_collection() -> chromadb.Collection:
    """Return (or create) the ChromaDB collection."""
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(
        path=str(CHROMA_DIR),
        settings=Settings(anonymized_telemetry=False),
    )
    return client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )


def ingest_document(doc_path: Path, collection: chromadb.Collection, model: SentenceTransformer):
    """Full pipeline for one document: extract → chunk → embed → upsert."""
    logger.info("Ingesting: %s", doc_path.name)
    ext = doc_path.suffix.lower()

    if ext == ".pdf":
        pages = extract_text_from_pdf(doc_path)
    elif ext in (".xlsx", ".xls"):
        pages = extract_text_from_excel(doc_path)
    elif ext == ".docx":
        pages = extract_text_from_word(doc_path)
    else:
        logger.warning("Unsupported file type: %s", ext)
        return 0

    if not pages:
        logger.warning("No text extracted from %s", doc_path.name)
        return 0

    chunks = chunk_pages(pages)
    if not chunks:
        return 0

    texts = [c["text"] for c in chunks]
    embeddings = model.encode(texts, batch_size=32, show_progress_bar=False).tolist()

    ids       = [f"{doc_path.stem}_p{c['page']}_c{c['chunk']}" for c in chunks]
    metadatas = [{"source": c["source"], "page": c["page"], "type": ext}  for c in chunks]

    # Upsert in batches of 100
    batch = 100
    for i in range(0, len(ids), batch):
        collection.upsert(
            ids        = ids[i:i+batch],
            embeddings = embeddings[i:i+batch],
            documents  = texts[i:i+batch],
            metadatas  = metadatas[i:i+batch],
        )

    logger.info("Ingested %d chunks from %s", len(chunks), doc_path.name)
    return len(chunks)


def run_ingestion(category: str | None = None) -> dict:
    """
    Ingest all documents (PDF, Excel, Word) from their respective directories.
    Returns summary {"total_files": int, "total_chunks": int, "pdf": int, "excel": int, "word": int}.
    """
    collection = get_chroma_collection()
    model      = get_embed_model()

    # Already-ingested document IDs (paginated to avoid SQLite variable limit)
    existing: set[str] = set()
    total = collection.count()
    if total > 0:
        batch_size = 5000
        for offset in range(0, total, batch_size):
            batch = collection.get(limit=batch_size, offset=offset)["ids"]
            existing.update(batch)

    # Collect files from all document directories
    doc_files: list[tuple[Path, str]] = []  # (path, type_label)

    pdf_root = PDF_DIR / category if category else PDF_DIR
    for f in pdf_root.rglob("*.pdf"):
        doc_files.append((f, "pdf"))

    excel_root = EXCEL_DIR / category if category else EXCEL_DIR
    for ext in ("*.xlsx", "*.xls"):
        for f in excel_root.rglob(ext):
            doc_files.append((f, "excel"))

    word_root = WORD_DIR / category if category else WORD_DIR
    for f in word_root.rglob("*.docx"):
        doc_files.append((f, "word"))

    logger.info("Found %d documents to process (PDF/Excel/Word)", len(doc_files))

    total_chunks = 0
    total_files  = 0
    type_counts  = {"pdf": 0, "excel": 0, "word": 0}

    for doc_path, doc_type in doc_files:
        # Skip if first chunk ID already exists
        first_id = f"{doc_path.stem}_p1_c0"
        if first_id in existing:
            logger.debug("Skipping already-ingested: %s", doc_path.name)
            continue

        chunks = ingest_document(doc_path, collection, model)
        if chunks:
            total_files  += 1
            total_chunks += chunks
            type_counts[doc_type] += 1

    summary = {
        "total_files":  total_files,
        "total_chunks": total_chunks,
        **type_counts,
    }
    logger.info("Ingestion complete: %s", summary)
    return summary


# ─── Retrieval ─────────────────────────────────────────────────────────────────
def retrieve_relevant_chunks(
    query: str,
    n_results: int = 5,
) -> List[dict]:
    """
    Embed query and retrieve top-n relevant chunks from ChromaDB.
    Returns list of {"text": str, "source": str, "page": int, "score": float}
    """
    model      = get_embed_model()
    collection = get_chroma_collection()

    if collection.count() == 0:
        logger.warning("ChromaDB collection is empty – run ingestion first.")
        return []

    query_embedding = model.encode([query]).tolist()
    results = collection.query(
        query_embeddings = query_embedding,
        n_results        = min(n_results, collection.count()),
        include          = ["documents", "metadatas", "distances"],
    )

    chunks = []
    for doc, meta, dist in zip(
        results["documents"][0],
        results["metadatas"][0],
        results["distances"][0],
    ):
        chunks.append({
            "text":   doc,
            "source": meta.get("source", "Unknown"),
            "page":   meta.get("page", 0),
            "score":  round(1 - dist, 4),  # cosine similarity
        })

    return chunks


if __name__ == "__main__":
    summary = run_ingestion()
    print("Ingestion summary:", summary)
